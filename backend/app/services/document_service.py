import json
import logging
import re
import threading
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

import fitz
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.logging_utils import log_event
from app.db.session import SessionLocal
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.schemas.documents import (
    DocumentActionResponse,
    DocumentBatchDeleteResponse,
    DocumentChunkDetailResponse,
    DocumentChunkItem,
    DocumentDetailResponse,
    DocumentItem,
    DocumentListResponse,
    SourcePageItem,
    UpdateDocumentRequest,
    UploadDocumentResponse,
)
from app.services.ocr_service import extract_text_with_ocr
from app.services.retrieval_service import current_embedding_model_name, generate_embedding, tokenize


HEADING_PATTERN = re.compile(r"^(#{1,6}\s+.+|第[一二三四五六七八九十百千万\d]+[章节部分条款].+)$")
SENTENCE_PATTERN = re.compile(r"(?<=[。！？；.!?;])")
SUPPORTED_FILE_TYPES = {"PDF", "DOCX", "TXT"}
logger = logging.getLogger(__name__)
PROCESSING_HINT = "文档处理中，暂不可查看完整内容。完成后会自动刷新状态。"
FAILED_HINT = "文档处理失败，请重新上传或删除后重试。"


# è¿åææ¡£æºæä»¶å­å¨ç®å½ï¼å¹¶ç¡®ä¿ç®å½å­å¨ã
def _storage_root() -> Path:
    root = Path(settings.document_storage_dir)
    if not root.is_absolute():
        root = Path(__file__).resolve().parents[2] / root
    root.mkdir(parents=True, exist_ok=True)
    return root


# æ ¹æ®æä»¶åååå®¹ç±»åç»ä¸å½ä¸åææ¡£ç±»åã
def _normalize_file_type(filename: str, content_type: str | None) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if content_type == "application/pdf" or suffix == "pdf":
        return "PDF"
    if content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    } or suffix == "docx":
        return "DOCX"
    return "TXT"


# ç»è®¡ææ¬åè¯åç token æ°éï¼ç¨äºåçåæ°æ®è®°å½ã
def _token_count(text: str) -> int:
    return len(tokenize(text))


# ç»ä¸ææ¬æ¢è¡æ ¼å¼å¹¶æè¡æåã
def _normalize_lines(text: str) -> list[str]:
    return [line.strip() for line in text.replace("\r\n", "\n").split("\n")]


# å°è§£æåºçæ®µè½åè¿½å å°ç»æå block åè¡¨ä¸­ã
def _append_block(blocks: list[dict[str, Any]], content: str, section_title: str | None, page_no: int | None) -> None:
    cleaned = content.strip()
    if not cleaned:
        return

    blocks.append(
        {
            "content": cleaned,
            "section_title": section_title,
            "page_no": page_no,
            "chunk_type": "section" if section_title and cleaned == section_title else "paragraph",
        }
    )


# ææ é¢åç©ºè¡æåæææç»æåæ®µè½åã
def _structured_blocks_from_text(text: str, page_no: int | None = None) -> list[dict[str, Any]]:
    lines = _normalize_lines(text)
    blocks: list[dict[str, Any]] = []
    buffer: list[str] = []
    current_section: str | None = None

    for line in lines:
        if not line:
            _append_block(blocks, "\n".join(buffer), current_section, page_no)
            buffer = []
            continue

        if HEADING_PATTERN.match(line) or (len(line) <= 28 and line.endswith(":")):
            _append_block(blocks, "\n".join(buffer), current_section, page_no)
            buffer = []
            current_section = line
            continue

        buffer.append(line)

    _append_block(blocks, "\n".join(buffer), current_section, page_no)
    return blocks


# å°è¿é¿æ®µè½æå¥å­æé¿åº¦åæéåæ£ç´¢çå°çæ®µã
def _split_long_text(text: str, max_chars: int = 360, overlap: int = 60) -> list[str]:
    cleaned = text.strip()
    if len(cleaned) <= max_chars:
        return [cleaned]

    sentences = [item.strip() for item in SENTENCE_PATTERN.split(cleaned) if item.strip()]
    if len(sentences) <= 1:
        step = max(max_chars - overlap, 80)
        sentences = [cleaned[index:index + max_chars] for index in range(0, len(cleaned), step)]

    pieces: list[str] = []
    current = ""
    for sentence in sentences:
        candidate = f"{current}{sentence}"
        if current and len(candidate) > max_chars:
            pieces.append(current.strip())
            tail = current[-overlap:] if overlap and len(current) > overlap else current
            current = f"{tail}{sentence}"
            continue
        current = candidate

    if current.strip():
        pieces.append(current.strip())

    return pieces or [cleaned]


# ä¼åä½¿ç¨ PyMuPDF æå PDF ææ¬ã
def _extract_pdf_with_pymupdf(content: bytes) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []
    with fitz.open(stream=content, filetype="pdf") as pdf:
        for index, page in enumerate(pdf, start=1):
            page_text = (page.get_text("text") or "").strip()
            if page_text:
                pages.append((index, page_text))
    return pages


# å½ PyMuPDF æ²¡æåå°ææ¬æ¶ï¼ä½¿ç¨ pypdf ååºè§£æ PDFã
def _extract_pdf_with_pypdf(content: bytes) -> list[tuple[int, str]]:
    reader = PdfReader(BytesIO(content))
    pages: list[tuple[int, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages.append((index, page_text))
    return pages


# æ ¹æ®ææ¡£ç±»åæååæãç»æååå PDF åé¡µåå®¹ã
def _extract_text_payload(
    content: bytes,
    filename: str,
    content_type: str | None,
) -> tuple[str, list[dict[str, Any]], list[dict[str, Any]]]:
    file_type = _normalize_file_type(filename, content_type)

    if file_type == "PDF":
        page_texts = _extract_pdf_with_pymupdf(content)
        if not page_texts:
            page_texts = _extract_pdf_with_pypdf(content)
        if not page_texts:
            page_texts = extract_text_with_ocr(content)

        if not page_texts:
            fallback = "当前 PDF 未提取到有效文本，可能是扫描件或图片型 PDF。"
            return fallback, [{"content": fallback, "section_title": "解析提示", "page_no": None, "chunk_type": "note"}], []

        raw_pages: list[str] = []
        blocks: list[dict[str, Any]] = []
        source_pages: list[dict[str, Any]] = []
        for page_no, page_text in page_texts:
            raw_pages.append(page_text)
            source_pages.append({"page_no": page_no, "content": page_text})
            blocks.extend(_structured_blocks_from_text(page_text, page_no=page_no))
        return "\n\n".join(raw_pages).strip(), blocks, source_pages

    if file_type == "DOCX":
        doc = DocxDocument(BytesIO(content))
        raw_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()).strip()
        return raw_text, _structured_blocks_from_text(raw_text), []

    raw_text = content.decode("utf-8", errors="ignore").strip()
    return raw_text, _structured_blocks_from_text(raw_text), []


# å°åæ block è½¬æ¢ä¸ºåç»­å»ºåºç¨çåçè½½è·ã
def _build_chunk_payloads(raw_text: str, blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    source_blocks = blocks or [
        {
            "content": raw_text.strip() or "未能从上传文档中提取有效文本内容。",
            "section_title": None,
            "page_no": None,
            "chunk_type": "paragraph",
        }
    ]
    chunk_payloads: list[dict[str, Any]] = []

    for block in source_blocks:
        for piece in _split_long_text(block["content"]):
            chunk_payloads.append(
                {
                    "content": piece,
                    "section_title": block.get("section_title"),
                    "page_no": block.get("page_no"),
                    "chunk_type": block.get("chunk_type") or "paragraph",
                    "token_count": _token_count(piece),
                }
            )

    return chunk_payloads


# éå»ºææ¡£çå¨é¨åçãåéåååé»æ¥å³ç³»ã
def _rebuild_chunks(db: Session, document: Document, chunk_payloads: list[dict[str, Any]]) -> None:
    db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()
    db.flush()

    created_chunks: list[DocumentChunk] = []
    embedding_model = current_embedding_model_name()
    for index, payload in enumerate(chunk_payloads):
        chunk = DocumentChunk(
            document_id=document.id,
            chunk_index=index,
            section_title=payload["section_title"],
            page_no=payload["page_no"],
            chunk_type=payload["chunk_type"],
            token_count=payload["token_count"],
            content=payload["content"],
            embedding_model=embedding_model,
            embedding=generate_embedding(payload["content"]),
        )
        db.add(chunk)
        db.flush()
        created_chunks.append(chunk)

    for index, chunk in enumerate(created_chunks):
        chunk.prev_chunk_id = created_chunks[index - 1].id if index > 0 else None
        chunk.next_chunk_id = created_chunks[index + 1].id if index < len(created_chunks) - 1 else None


# å°ä¸ä¼ çåå§æä»¶æä¹åå°æ¬å°å­å¨ç®å½ã
def _save_source_file(document_id: str, filename: str, content: bytes) -> Path:
    storage_root = _storage_root()
    target_dir = storage_root / document_id
    target_dir.mkdir(parents=True, exist_ok=True)
    safe_name = Path(filename).name or f"{uuid4()}.bin"
    target_path = target_dir / safe_name
    target_path.write_bytes(content)
    return target_path


# å é¤ææ¡£å¯¹åºçæºæä»¶åç©ºç®å½ã
def _delete_source_file(document: Document) -> None:
    if not document.source_file_path:
        return

    path = Path(document.source_file_path)
    try:
        if path.exists():
            path.unlink()
        if path.parent.exists() and not any(path.parent.iterdir()):
            path.parent.rmdir()
    except OSError:
        return


# å¤æ­ææ¡£æºæä»¶å½åæ¯å¦ä»ç¶å­å¨ã
def _document_source_exists(document: Document) -> bool:
    return bool(document.source_file_path and Path(document.source_file_path).exists())


# å°æ°æ®åºéçåé¡µ JSON ååºååä¸ºæ¥å£è¾åºç»æã
def _serialize_source_pages(document: Document) -> list[SourcePageItem]:
    if not document.source_pages_json:
        return []
    try:
        data = json.loads(document.source_pages_json)
    except json.JSONDecodeError:
        return []
    return [SourcePageItem.model_validate(item) for item in data]


# å°ææ¡£æ¨¡åå¯¹è±¡è½¬æ¢ä¸ºåè¡¨å±ç¤ºç¨çååºç»æã
def _serialize_document_item(document: Document) -> DocumentItem:
    return DocumentItem(
        id=document.id,
        title=document.title,
        filename=document.filename,
        file_type=document.file_type or _normalize_file_type(document.filename, document.content_type),
        source_file_path=document.source_file_path,
        source_file_size=document.source_file_size,
        source_file_exists=_document_source_exists(document),
        status=document.status,
        summary=document.summary,
        chunk_count=document.chunk_count,
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


# æ ¡éªæ¬æ¬¡ä¸ä¼ æ é¢æ¯å¦éå¤ï¼å¹¶æ¦æªå·²å­å¨çååææ¡£ã
def ensure_unique_titles(
    db: Session,
    titles: list[str],
    owner_id: str | None = None,
) -> None:
    normalized = [title.strip() for title in titles if title and title.strip()]
    duplicates_in_request = {name for name in normalized if normalized.count(name) > 1}
    if duplicates_in_request:
        duplicate_title = sorted(duplicates_in_request)[0]
        raise ValueError(f"document title duplicated in request: {duplicate_title}")

    for title in normalized:
        query = db.query(Document).filter(Document.title == title)
        if owner_id:
            query = query.filter(Document.owner_id == owner_id)
        existing = query.first()
        if existing:
            raise ValueError(f"document title already exists: {title}")


# ååå»ºå¤çä¸­ç¶æçå ä½ææ¡£è®°å½ï¼å¹¶ä¿å­æºæä»¶ã
def _create_document_placeholder(
    db: Session,
    title: str,
    filename: str,
    content_type: str | None,
    content: bytes,
    owner_id: str | None = None,
) -> UploadDocumentResponse:
    safe_filename = Path(filename).name or "uploaded.bin"
    now = datetime.utcnow()
    document = Document(
        owner_id=owner_id,
        title=title.strip(),
        filename=safe_filename,
        content_type=content_type,
        file_type=_normalize_file_type(safe_filename, content_type),
        status="processing",
        summary=PROCESSING_HINT,
        source_text="",
        source_pages_json=None,
        chunk_count=0,
        updated_at=now,
    )
    db.add(document)
    db.flush()

    source_path = _save_source_file(document.id, safe_filename, content)
    document.source_file_path = str(source_path)
    document.source_file_size = len(content)
    db.commit()
    db.refresh(document)

    return UploadDocumentResponse(
        id=document.id,
        title=document.title,
        status=document.status,
        chunk_count=document.chunk_count,
    )


# å¨åå°çº¿ç¨ä¸­å®æè§£æãåçãåéååç¶ææ´æ°ã
def _process_document_placeholder(
    document_id: str,
    filename: str,
    content_type: str | None,
    content: bytes,
) -> None:
    started = datetime.utcnow()
    safe_filename = Path(filename).name or "uploaded.bin"
    db = SessionLocal()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            log_event(
                logger,
                "document.create.missing_placeholder",
                level=logging.ERROR,
                document_id=document_id,
                filename=safe_filename,
            )
            return

        raw_text, blocks, source_pages = _extract_text_payload(content, safe_filename, content_type)
        chunk_payloads = _build_chunk_payloads(raw_text, blocks)
        document.summary = raw_text[:220] if raw_text else "No valid text extracted."
        document.source_text = raw_text
        document.source_pages_json = json.dumps(source_pages, ensure_ascii=False) if source_pages else None
        document.chunk_count = len(chunk_payloads)
        _rebuild_chunks(db, document, chunk_payloads)
        document.status = "indexed"
        document.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(document)
        elapsed_ms = round((datetime.utcnow() - started).total_seconds() * 1000, 2)
        log_event(
            logger,
            "document.create.completed",
            document_id=document.id,
            title=document.title,
            filename=document.filename,
            file_type=document.file_type,
            chunk_count=document.chunk_count,
            elapsed_ms=elapsed_ms,
        )
    except Exception as exc:
        db.rollback()
        failed_document = db.query(Document).filter(Document.id == document_id).first()
        if failed_document:
            failed_document.status = "failed"
            failed_document.summary = FAILED_HINT
            failed_document.source_text = ""
            failed_document.source_pages_json = None
            failed_document.chunk_count = 0
            failed_document.updated_at = datetime.utcnow()
            db.commit()
        log_event(
            logger,
            "document.create.failed",
            level=logging.ERROR,
            document_id=document_id,
            filename=safe_filename,
            error=repr(exc),
        )
    finally:
        db.close()


# åå»ºææ¡£å ä½è®°å½ï¼å¹¶å¼æ­¥å¯å¨åå°å¤çæµç¨ã
def create_document(
    db: Session,
    title: str,
    filename: str,
    content_type: str | None,
    content: bytes,
    owner_id: str | None = None,
) -> UploadDocumentResponse:
    result = _create_document_placeholder(
        db=db,
        title=title,
        filename=filename,
        content_type=content_type,
        content=content,
        owner_id=owner_id,
    )
    worker = threading.Thread(
        target=_process_document_placeholder,
        args=(result.id, filename, content_type, content),
        daemon=True,
    )
    worker.start()
    return result


# è¿åç¬¦åç­éæ¡ä»¶çææ¡£åè¡¨ã
def list_documents(db: Session, query: str | None = None, file_type: str | None = None) -> DocumentListResponse:
    documents_query = db.query(Document)
    if query:
        like = f"%{query.strip()}%"
        documents_query = documents_query.filter(
            (Document.title.ilike(like)) | (Document.filename.ilike(like)) | (Document.summary.ilike(like))
        )
    if file_type and file_type.upper() in SUPPORTED_FILE_TYPES:
        documents_query = documents_query.filter(Document.file_type == file_type.upper())

    items = documents_query.order_by(Document.updated_at.desc(), Document.created_at.desc()).all()
    log_event(logger, "document.list.completed", query=query or "", file_type=file_type or "", count=len(items))
    return DocumentListResponse(items=[_serialize_document_item(item) for item in items])


# è¿ååä¸ªææ¡£çå®æ´è¯¦æãåæååçåè¡¨ã
def get_document(db: Session, document_id: str) -> DocumentDetailResponse | None:
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        log_event(logger, "document.detail.missing", level=logging.WARNING, document_id=document_id)
        return None

    item = _serialize_document_item(document)
    return DocumentDetailResponse(
        **item.model_dump(),
        content_type=document.content_type,
        source_text=document.source_text or (PROCESSING_HINT if document.status != "indexed" else ""),
        source_pages=_serialize_source_pages(document),
        chunks=[
            DocumentChunkItem(
                id=chunk.id,
                chunk_index=chunk.chunk_index,
                section_title=chunk.section_title,
                page_no=chunk.page_no,
                chunk_type=chunk.chunk_type,
                token_count=chunk.token_count,
                content=chunk.content,
            )
            for chunk in sorted(document.chunks, key=lambda entry: entry.chunk_index)
        ],
    )


# è¿ååä¸ªåççä¸ä¸æä¿¡æ¯ãé»æ¥åçååæå®ä½åå®¹ã
def get_chunk_detail(db: Session, chunk_id: str) -> DocumentChunkDetailResponse | None:
    chunk = db.query(DocumentChunk).filter(DocumentChunk.id == chunk_id).first()
    if not chunk or not chunk.document:
        log_event(logger, "document.chunk_detail.missing", level=logging.WARNING, chunk_id=chunk_id)
        return None

    document = chunk.document
    source_pages = _serialize_source_pages(document)
    source_page_content = None
    if chunk.page_no:
        source_page_content = next((item.content for item in source_pages if item.page_no == chunk.page_no), None)

    previous_chunk = None
    if chunk.prev_chunk_id:
        prev = db.query(DocumentChunk).filter(DocumentChunk.id == chunk.prev_chunk_id).first()
        if prev:
            previous_chunk = {
                "id": prev.id,
                "chunk_index": prev.chunk_index,
                "section_title": prev.section_title,
                "page_no": prev.page_no,
                "content": prev.content,
            }

    next_chunk = None
    if chunk.next_chunk_id:
        nxt = db.query(DocumentChunk).filter(DocumentChunk.id == chunk.next_chunk_id).first()
        if nxt:
            next_chunk = {
                "id": nxt.id,
                "chunk_index": nxt.chunk_index,
                "section_title": nxt.section_title,
                "page_no": nxt.page_no,
                "content": nxt.content,
            }

    return DocumentChunkDetailResponse(
        chunk_id=chunk.id,
        document_id=document.id,
        document_title=document.title,
        filename=document.filename,
        file_type=document.file_type or _normalize_file_type(document.filename, document.content_type),
        content_type=document.content_type,
        page_no=chunk.page_no,
        section_title=chunk.section_title,
        chunk_index=chunk.chunk_index,
        snippet=chunk.content[:320],
        content=chunk.content,
        previous_chunk=previous_chunk,
        next_chunk=next_chunk,
        source_text=document.source_text,
        source_page_content=source_page_content,
        source_file_exists=_document_source_exists(document),
    )


# æ´æ°ææ¡£æ é¢ææè¦ï¼å¹¶è¿åææ°è¯¦æã
def update_document(db: Session, document_id: str, payload: UpdateDocumentRequest) -> DocumentDetailResponse | None:
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        return None

    if payload.title is not None:
        document.title = payload.title.strip() or document.title
    if payload.summary is not None:
        document.summary = payload.summary.strip()
    document.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(document)
    return get_document(db, document_id)


# å é¤åä¸ªææ¡£åå¶çº§èæ°æ®ã
def delete_document(db: Session, document_id: str) -> DocumentActionResponse | None:
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        log_event(logger, "document.delete.missing", level=logging.WARNING, document_id=document_id)
        return None

    _delete_source_file(document)
    db.delete(document)
    db.commit()
    log_event(logger, "document.delete.completed", document_id=document_id, filename=document.filename)
    return DocumentActionResponse(id=document_id, message="Document deleted")


# æ¹éå é¤å¤ä¸ªææ¡£å¹¶è¿åå é¤ç»ææè¦ã
def batch_delete_documents(db: Session, document_ids: list[str]) -> DocumentBatchDeleteResponse:
    normalized_ids = [item for item in dict.fromkeys(document_ids) if item]
    documents = db.query(Document).filter(Document.id.in_(normalized_ids)).all() if normalized_ids else []

    deleted_ids: list[str] = []
    for document in documents:
        _delete_source_file(document)
        deleted_ids.append(document.id)
        db.delete(document)

    db.commit()
    log_event(logger, "document.batch_delete.completed", requested=len(normalized_ids), deleted_count=len(deleted_ids))
    return DocumentBatchDeleteResponse(
        ids=deleted_ids,
        deleted_count=len(deleted_ids),
        message=f"Deleted {len(deleted_ids)} documents",
    )


# è¿åå¯ä¸è½½æºæä»¶çææ¡£å¯¹è±¡åæä»¶è·¯å¾ã
def get_source_file_path(db: Session, document_id: str) -> tuple[Document, Path] | None:
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document or not document.source_file_path:
        return None

    path = Path(document.source_file_path)
    if not path.exists():
        return None
    return document, path
